"""
country_profiles_extract.py

Reads the two source .docx files in country_profiles_source/ and produces
server/data/country_profiles.json, keyed by the same zero-padded 3-digit
ISO numeric country code country_data.json already uses:

    {
      "<iso_numeric>": {
        "name": "<country name>",
        "overview": {
          "paragraphs": ["...", "..."],
          "references": ["<APA citation>", ...]
        },
        "dashboard_note": {                 # null if this country has no
          "paragraphs": ["...", "..."],     # entry in the second source
          "references": ["<citation>", ...] # doc (see below)
        } | null
      },
      ...
    }

Source documents (data_scripts/country_profiles_source/):

  1. country_profile_section_of_our_international_observatory_website.docx
     "Country Trauma Profiles" - the full survey, one entry per country:
       **<Country>**
       <one or more paragraphs>
       **Reference:** <APA citation>
       **Reference:** <APA citation>            # optional, one or more

  2. country_profiles_clean_no_spreadsheet_citations.docx
     "One-Page Country Collective Trauma Profiles" - a smaller companion
     set (countries with GTBI/ETTI dashboard data) tying the same
     history to the Observatory's dashboard scores:
       **<Country>**
       <one or more paragraphs>
       **APA 7 references**
       <citation 1>
       <citation 2>
       ...
     Plus a trailing **General reference** section (the shared
     International IDEA citation used in this document's own intro,
     not any single country's in-text reference) - intentionally
     dropped rather than attached to every country.

Every country in source #1 also appearing in source #2 gets a
"dashboard_note" merged in; country names are resolved to ISO numeric
codes via server/data/country_data.json's own "name" fields (with a
small alias table for the handful that don't match verbatim - old vs.
current official names, "Turkey" vs "Türkiye", etc.). Kosovo has no
ISO 3166-1 numeric code at all and is dropped from the output for that
reason - its profile still exists in the source docx if it's ever
needed, it just has nowhere to live in this code-keyed format.

Usage:
    python3 country_profiles_extract.py
    python3 country_profiles_extract.py --country-data ../data/country_data.json -o ../data/country_profiles.json
"""
import argparse
import json
import os
import re

from docx import Document

DEFAULT_SOURCE_DIR = os.path.join(os.path.dirname(__file__), "country_profiles_source")
DEFAULT_SURVEY_DOCX = "country_profile_section_of_our_international_observatory_website.docx"
DEFAULT_DASHBOARD_DOCX = "country_profiles_clean_no_spreadsheet_citations.docx"

SURVEY_TITLE = "Country Trauma Profiles"
DASHBOARD_TITLE = "One-Page Country Collective Trauma Profiles"
DASHBOARD_TRAILING_SECTION = "General reference"  # shared citation, not per-country - dropped

# Country names in the source docx that don't match country_data.json's
# "name" field verbatim (older/alternate English names vs. the current
# ISO 3166-1 / pycountry official short names).
NAME_ALIASES = {
    "Bolivia": "Bolivia, Plurinational State of",
    "Congo (Republic of the Congo)": "Republic of the Congo",
    "Democratic Republic of the Congo": None,  # resolved by numeric code directly, see below
    "Laos": "Lao People's Democratic Republic",
    "Micronesia": "Micronesia, Federated States of",
    "Moldova": "Moldova, Republic of",
    "North Korea": "Korea, Democratic People's Republic of",
    "South Korea": "Korea, Republic of",
    "Palestine": "Palestine, State of",
    "Russia": "Russian Federation",
    "Syria": "Syrian Arab Republic",
    "Taiwan": "Taiwan, Province of China",
    "Tanzania": "Tanzania, United Republic of",
    "Turkey": "Türkiye",
    "Vatican City": "Holy See (Vatican City State)",
    "Venezuela": "Venezuela, Bolivarian Republic of",
    "Vietnam": "Viet Nam",
}
# ISO numeric code for the one alias above that isn't a simple name
# swap - country_data.json's own "Congo, The Democratic Republic of the"
# → "Congo" name collision (a pre-existing display-name quirk shared with
# the Republic of the Congo, both currently labelled just "Congo") means
# matching by name alone is ambiguous here, so this one is hardcoded.
DRC_NUMERIC_CODE = "180"


def is_heading(paragraph):
    """A heading in these docs is a paragraph that's entirely bold text."""
    text = paragraph.text.strip()
    if not text or not paragraph.runs:
        return None
    if all(run.bold for run in paragraph.runs if run.text.strip()):
        return text
    return None


def parse_survey_docx(path):
    """Source #1: **Country** / paragraph(s) / one or more **Reference:** <citation> lines"""
    doc = Document(path)
    countries = {}
    current = None
    paragraphs = []
    references = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        heading = is_heading(p)
        if heading and heading != SURVEY_TITLE:
            if current:
                countries[current] = {"paragraphs": paragraphs, "references": references}
            current = heading
            paragraphs = []
            references = []
            continue

        if current is None:
            continue  # still in the intro paragraph(s) before the first country

        ref_match = re.match(r"^Reference:\s*(.+)$", text)
        if ref_match:
            # A country can have more than one - each is its own
            # "Reference:"-prefixed paragraph in the source doc, appended
            # here in document order rather than overwriting a single
            # value. Almost every country still has exactly one, so this
            # produces a one-item list for those - same shape either way.
            references.append(ref_match.group(1).strip())
        else:
            paragraphs.append(text)

    if current:
        countries[current] = {"paragraphs": paragraphs, "references": references}
    return countries


def parse_dashboard_docx(path):
    """Source #2: **Country** / paragraph(s) / **APA 7 references** / citation lines"""
    doc = Document(path)
    countries = {}
    current = None
    paragraphs = []
    references = []
    in_refs = False

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        heading = is_heading(p)
        if heading and heading != DASHBOARD_TITLE:
            if current and current != DASHBOARD_TRAILING_SECTION:
                countries[current] = {"paragraphs": paragraphs, "references": references}
            current = heading
            paragraphs = []
            references = []
            in_refs = False
            continue

        if current is None:
            continue

        if text.startswith("APA 7 references"):
            in_refs = True
            continue

        if in_refs:
            references.append(text)
        else:
            paragraphs.append(text)

    if current and current != DASHBOARD_TRAILING_SECTION:
        countries[current] = {"paragraphs": paragraphs, "references": references}
    return countries


def resolve_code(name, name_to_code):
    if name in name_to_code:
        return name_to_code[name]
    if name == "Democratic Republic of the Congo":
        return DRC_NUMERIC_CODE
    alias = NAME_ALIASES.get(name)
    if alias and alias in name_to_code:
        return name_to_code[alias]
    return None


def build_profiles(survey, dashboard, country_data):
    name_to_code = {v["name"]: k for k, v in country_data.items()}
    profiles = {}
    skipped = []

    for name, entry in survey.items():
        code = resolve_code(name, name_to_code)
        if code is None:
            skipped.append(name)
            continue

        dashboard_entry = dashboard.get(name)
        profiles[code] = {
            "name": name,
            "overview": {
                "paragraphs": entry["paragraphs"],
                "references": entry["references"],
            },
            "dashboard_note": (
                {
                    "paragraphs": dashboard_entry["paragraphs"],
                    "references": dashboard_entry["references"],
                }
                if dashboard_entry
                else None
            ),
        }

    return profiles, skipped


def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument(
        "--source-dir", default=DEFAULT_SOURCE_DIR,
        help="Directory containing the two source .docx files (default: country_profiles_source/ next to this script)",
    )
    parser.add_argument(
        "--country-data", default=os.path.join(os.path.dirname(__file__), "..", "data", "country_data.json"),
        help="Path to country_data.json, used to resolve country names to ISO numeric codes",
    )
    parser.add_argument(
        "-o", "--output", default=os.path.join(os.path.dirname(__file__), "..", "data", "country_profiles.json"),
        help="Output path for country_profiles.json",
    )
    args = parser.parse_args()

    survey_path = os.path.join(args.source_dir, DEFAULT_SURVEY_DOCX)
    dashboard_path = os.path.join(args.source_dir, DEFAULT_DASHBOARD_DOCX)

    survey = parse_survey_docx(survey_path)
    dashboard = parse_dashboard_docx(dashboard_path)

    with open(args.country_data, encoding="utf-8") as f:
        country_data = json.load(f)

    profiles, skipped = build_profiles(survey, dashboard, country_data)

    with open(args.output, "w", encoding="utf-8") as f:
        json.dump(profiles, f, ensure_ascii=False, indent=2)

    with_dashboard = sum(1 for v in profiles.values() if v["dashboard_note"])
    print(f"Wrote {len(profiles)} country profiles to {args.output}")
    print(f"  {with_dashboard} include an Observatory dashboard note")
    if skipped:
        print(f"  Skipped (no ISO numeric code match): {', '.join(skipped)}")


if __name__ == "__main__":
    main()