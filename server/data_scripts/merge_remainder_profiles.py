"""
merge_remainder_profiles.py

One-off utility: merges a "remainder" country-profile docx (entries for
countries missing from the canonical survey source) into the existing
canonical survey docx, producing a single combined file in the same
**Country** / paragraph(s) / **Reference:** <citation> format
country_profiles.extract.py's parse_survey_docx() expects.

Why this exists rather than just uploading the remainder file through
the admin Control panel's docx upload: that flow (see
country_profiles_upload.py) always treats an uploaded "survey" file as
a full replacement of the current canonical source, not a merge - it
has no notion of "these are just the missing entries, keep the
existing 165 too." Since a remainder file by definition only covers
what the current canonical file is missing, replacing outright would
lose those 165 entries. This script does the merge once, up front, so
the *result* can be uploaded/committed as the new canonical file and
the normal single-file replace flow works fine for any future update.

Usage (from server/data_scripts/):
    python3 merge_remainder_profiles.py <path-to-remainder.docx>

Effects:
  1. Parses both the current canonical survey docx and the remainder
     docx with parse_survey_docx() (skipping any all-bold "PART n: ..."
     section-divider paragraphs in the remainder file, which aren't
     country entries).
  2. Errors out if any country name appears in both - a remainder file
     should only ever contain entries the canonical file is missing;
     an overlap most likely means this has already been merged once,
     or the wrong file was passed.
  3. Moves the current canonical file into country_profiles_source/old/,
     timestamped (same rotate pattern country_profiles_upload.py's
     rotate_source_docx() uses for admin uploads), then writes the
     merged set (all countries, name-sorted) to the canonical filename.
  4. Re-runs the normal extract -> server/data/country_profiles.json
     build, same as running country_profiles.extract.py directly.
"""

import datetime
import os
import shutil
import sys
import unicodedata

from docx import Document

import importlib.util

_HERE = os.path.dirname(os.path.abspath(__file__))
_EXTRACT_SCRIPT_PATH = os.path.join(_HERE, "country_profiles_extract.py")
_spec = importlib.util.spec_from_file_location("country_profiles_extract", _EXTRACT_SCRIPT_PATH)
cpe = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(cpe)

SOURCE_DIR = cpe.DEFAULT_SOURCE_DIR
CANONICAL_PATH = os.path.join(SOURCE_DIR, cpe.DEFAULT_SURVEY_DOCX)

TITLE = "Country Trauma Profiles"
INTRO = (
    "This section presents a comprehensive, alphabetically organized survey of collective trauma across the "
    "nations of the world, examining each country's most significant historical wounds\u2014including genocide, "
    "war, oppression, epidemic, and disaster\u2014alongside how these legacies continue to shape human suffering "
    "and psychiatric aftermath in the present day."
)


def _sort_key(name):
    """Case-insensitive, accent-insensitive sort key (purely cosmetic -
    the app's own country grouping/sorting happens client-side from
    country_data.json's names, not from this file's paragraph order)."""
    stripped = "".join(c for c in unicodedata.normalize("NFKD", name) if not unicodedata.combining(c))
    return stripped.lower()


def parse_remainder(path):
    """Same format as the canonical survey file, plus bold "PART n: ..."
    section dividers that aren't real country entries and get dropped."""
    parsed = cpe.parse_survey_docx(path)
    return {name: entry for name, entry in parsed.items() if not name.startswith("PART ")}


def build_merged_docx(entries, output_path):
    doc = Document()
    title_p = doc.add_paragraph()
    title_p.add_run(TITLE).bold = True
    doc.add_paragraph(INTRO)
    doc.add_paragraph("")
    doc.add_paragraph("")

    for name in sorted(entries, key=_sort_key):
        entry = entries[name]
        heading_p = doc.add_paragraph()
        heading_p.add_run(name).bold = True
        for para_text in entry["paragraphs"]:
            doc.add_paragraph(para_text)
        ref_p = doc.add_paragraph()
        ref_p.add_run("Reference:").bold = True
        ref_p.add_run(f" {entry['reference']}")

    doc.save(output_path)


def main():
    if len(sys.argv) != 2:
        print(f"Usage: python3 {os.path.basename(__file__)} <path-to-remainder.docx>")
        sys.exit(1)

    remainder_path = sys.argv[1]

    existing = cpe.parse_survey_docx(CANONICAL_PATH)
    remainder = parse_remainder(remainder_path)

    overlap = sorted(set(existing) & set(remainder))
    if overlap:
        print(f"Error: {len(overlap)} name(s) already exist in the canonical file - refusing to merge: {overlap}")
        sys.exit(1)

    merged = {**existing, **remainder}
    print(f"Merging {len(remainder)} new entries into {len(existing)} existing -> {len(merged)} total")

    old_dir = os.path.join(SOURCE_DIR, "old")
    os.makedirs(old_dir, exist_ok=True)
    if os.path.isfile(CANONICAL_PATH):
        timestamp = datetime.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
        backup_path = os.path.join(old_dir, f"{timestamp}_{cpe.DEFAULT_SURVEY_DOCX}")
        shutil.copy2(CANONICAL_PATH, backup_path)
        print(f"Backed up previous canonical file to {backup_path}")

    build_merged_docx(merged, CANONICAL_PATH)
    print(f"Wrote merged survey docx to {CANONICAL_PATH}")

    # Re-run the normal extract -> country_profiles.json build.
    import subprocess
    subprocess.run([sys.executable, _EXTRACT_SCRIPT_PATH], cwd=_HERE, check=True)


if __name__ == "__main__":
    main()